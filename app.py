
# --- 1. Impor Library ---
import streamlit as st
from langchain_groq import ChatGroq
from langchain.prompts import PromptTemplate
import PyPDF2
from docx import Document as DocxDocument
import requests
from bs4 import BeautifulSoup
import re
import os
import base64

# --- 2. Konfigurasi Halaman & Desain (CSS) ---
# Membaca ikon halaman dari file lokal
st.set_page_config(page_title="SMARTPAPER.AI", layout="wide", page_icon="SMAPER.png")

# Fungsi untuk mengubah gambar menjadi base64 untuk ditampilkan di HTML
def get_base64_of_bin_file(bin_file):
    if os.path.exists(bin_file):
        with open(bin_file, 'rb') as f:
            data = f.read()
        return base64.b64encode(data).decode()
    return ""

# CSS Kustom untuk tampilan dan nuansa aplikasi
CUSTOM_CSS = """
<style>
    @import url('https://fonts.googleapis.com/css2?family=Poppins:wght@300;400;600;700&display=swap');
    
    html, body, [class*="st-"], .st-emotion-cache-1v0mbdj {
        font-family: 'Poppins', sans-serif;
        color: #000000 !important; /* TEKS HITAM TEGAS */
    }

    /* --- BACKGROUND & LAYOUT --- */
    .stApp {
        background-color: #FFFFFF; /* Latar belakang putih bersih */
    }
    
    .block-container {
        padding: 2rem 5rem !important;
    }
    
    /* --- HEADER / NAVBAR --- */
    .header {
        display: flex;
        justify-content: space-between;
        align-items: center;
        padding-bottom: 1rem;
        border-bottom: 1px solid #dee2e6;
        margin-bottom: 4rem;
    }
    .header .logo { display: flex; align-items: center; gap: 15px; }
    .header .logo img { width: 50px; height: 50px; } /* Ukuran logo diperbesar */
    .header .logo .title { font-size: 2rem; font-weight: 700; color: #000000; }
    
    /* --- TYPOGRAPHY --- */
    h1 {
        font-size: 3.8rem;
        font-weight: 700;
        color: #000000;
        line-height: 1.2;
        margin-bottom: 1rem;
    }
    .tagline {
        font-size: 1.2rem;
        font-weight: 300;
        color: #495057;
        margin-bottom: 3rem;
    }
    
    /* --- INPUT FIELDS --- */
    .stTextInput label {
        font-weight: 600;
        color: #000000 !important;
    }
    .stTextInput > div > div > input {
        background-color: #FFFFFF !important; /* Latar belakang input PUTIH */
        color: #000000 !important; /* Teks input HITAM agar terlihat */
        border: 1px solid #ced4da;
        border-radius: 8px;
        padding: 12px 15px;
        transition: all 0.2s ease-in-out;
    }
    .stTextInput > div > div > input:focus {
        border-color: #009EDB; /* Warna border biru saat aktif */
        box-shadow: 0 0 0 3px rgba(0, 158, 219, 0.15);
    }

    /* --- BUTTONS --- */
    .stButton>button {
        font-weight: 600;
        border-radius: 8px;
        padding: 14px 30px;
        transition: all 0.2s ease;
        border: none;
        background-color: #009EDB; /* UN Blue */
        color: white;
    }
    .stButton>button:hover {
        background-color: #007fad;
        transform: translateY(-2px);
        box-shadow: 0 4px 10px rgba(0, 158, 219, 0.3);
    }

    /* --- RIGHT COLUMN IMAGE --- */
    [data-testid="column"]:has(img) img {
        border-radius: 12px;
        box-shadow: 0 10px 30px -10px rgba(0,0,0,0.2);
    }

    /* --- GAYA UNTUK RADIO BUTTON & FILE UPLOADER --- */
    /* Radio Buttons */
    .stRadio > div {
        flex-direction: row;
        gap: 10px;
    }
    .stRadio label {
        display: flex;
        justify-content: center;
        align-items: center;
        padding: 8px 20px;
        border: 1px solid #dee2e6;
        border-radius: 8px;
        transition: all 0.2s ease;
        cursor: pointer;
        background-color: #f8f9fa; /* Warna default abu-abu muda */
    }
    .stRadio > div > label:has(input:checked) {
        background-color: #009EDB; /* Warna biru saat dipilih */
        color: white !important;
        border-color: #009EDB;
    }
    .stRadio > div > label:not(:has(input:checked)):hover {
        background-color: #e9ecef; /* Efek hover */
        border-color: #adb5bd;
    }

    /* File Uploader */
    [data-testid="stFileUploader"] {
        border: 2px dashed #009EDB;
        background-color: rgba(0, 158, 219, 0.05); /* Warna biru sangat terang */
        border-radius: 12px;
        padding: 2rem;
    }
    [data-testid="stFileUploader"] p, [data-testid="stFileUploader"] small {
        color: #000000;
    }
    [data-testid="stFileUploader"] button {
        background-color: #343a40;
        color: white;
        border: none;
        border-radius: 8px;
        padding: 8px 16px;
        transition: all 0.2s ease;
    }
    [data-testid="stFileUploader"] button:hover {
        background-color: #495057;
    }
</style>
"""
st.markdown(CUSTOM_CSS, unsafe_allow_html=True)

# --- Fungsi Inti ---
@st.cache_resource
def get_llm():
    # --- PERINGATAN KEAMANAN ---
    # Kunci API dimasukkan langsung di sini sesuai permintaan untuk fungsionalitas segera.
    # SANGAT DISARANKAN untuk memindahkan kunci ini ke Streamlit Secrets
    # sebelum membagikan atau men-deploy aplikasi ini untuk menghindari penyalahgunaan.
    # Cara yang aman: groq_api_key = st.secrets["GROQ_API_KEY"]
    try:
        groq_api_key = "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"
        if not groq_api_key:
            st.error("Kunci API Groq tidak disediakan. Aplikasi tidak dapat berfungsi.")
            return None
        return ChatGroq(temperature=0, model_name="llama3-8b-8192", api_key=groq_api_key)
    except Exception as e:
        # Menangkap error lain, misalnya jika kunci tidak valid.
        st.error(f"Gagal memuat model AI. Pastikan API Key Anda valid. Error: {e}")
        return None


def extract_text(source, source_type):
    if source_type == 'file':
        if source.type == "application/pdf":
            try:
                reader = PyPDF2.PdfReader(source)
                return [page.extract_text() for page in reader.pages if page.extract_text()]
            except Exception as e: return [f"Error PDF: {e}"]
        elif source.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            try:
                doc = DocxDocument(source)
                return ["\n".join([para.text for para in doc.paragraphs])]
            except Exception as e: return [f"Error DOCX: {e}"]
    elif source_type == 'url':
        try:
            # Mengubah link Google Drive menjadi link download langsung jika terdeteksi
            if "drive.google.com" in source:
                file_id = source.split('/d/')[1].split('/')[0]
                source = f'https://drive.google.com/uc?export=download&id={file_id}'

            headers = {'User-Agent': 'Mozilla/5.0'}
            response = requests.get(source, headers=headers, timeout=20)
            response.raise_for_status()
            
            # Jika sumbernya adalah URL Google Drive, kita perlu menyimpannya sementara
            # karena PyPDF2/python-docx tidak bisa membaca langsung dari stream respons web
            temp_file_path = "temp_downloaded_file"
            with open(temp_file_path, "wb") as f:
                f.write(response.content)
            
            # Cek tipe file dari header atau coba tebak dari URL
            content_type = response.headers.get('content-type', '').lower()
            if 'pdf' in content_type or temp_file_path.endswith('.pdf'):
                 with open(temp_file_path, "rb") as f:
                    reader = PyPDF2.PdfReader(f)
                    return [page.extract_text() for page in reader.pages if page.extract_text()]
            elif 'word' in content_type or temp_file_path.endswith('.docx'):
                 with open(temp_file_path, "rb") as f:
                    doc = DocxDocument(f)
                    return ["\n".join([para.text for para in doc.paragraphs])]
            else: # Fallback ke BeautifulSoup untuk URL non-dokumen
                soup = BeautifulSoup(response.content, 'html.parser')
                for element in soup(["script", "style", "nav", "footer", "header"]): element.decompose()
                return [soup.get_text(separator='\n', strip=True)]

        except requests.RequestException as e: return [f"Error URL: {e}"]
        finally:
            if os.path.exists("temp_downloaded_file"):
                os.remove("temp_downloaded_file") # Hapus file sementara
    return None

def generate_summary(text_content, page_num=None):
    llm = get_llm()
    if not llm or not text_content: return "Gagal membuat ringkasan."
    
    title = f"halaman {page_num}" if page_num else "keseluruhan dokumen"
    prompt = PromptTemplate.from_template(
        f"Anda adalah analis riset AI. Buat ringkasan komprehensif dari teks {title} berikut dalam format poin-poin (bullet points) dalam Bahasa Indonesia. Fokus pada ide utama dan poin-poin kunci.\n\nTeks:\n---\n{{text}}\n---"
    )
    chain = prompt | llm
    summary = chain.invoke({"text": text_content[:12000]}).content
    return summary

def answer_question(document_pages, question):
    llm = get_llm()
    if not llm: return "Model AI tidak tersedia."
    
    page_match = re.search(r'(halaman|hlm|page|halkamn)[\s:]*(\d+)', question.lower())
    
    if page_match:
        page_number = int(page_match.group(2))
        page_index = page_number - 1
        
        if 0 <= page_index < len(document_pages):
            page_content = document_pages[page_index]
            summary = generate_summary(page_content, page_num=page_number)
            return f"**Ringkasan untuk Halaman {page_number}:**\n\n" + summary
        else:
            return f"Maaf, dokumen ini hanya memiliki {len(document_pages)} halaman. Tidak dapat menemukan halaman {page_number}."
    else:
        full_context = "\n".join(document_pages)
        prompt = PromptTemplate.from_template(
            "Berdasarkan konteks dokumen berikut, jawablah pertanyaan pengguna dengan jelas dalam Bahasa Indonesia.\n\nKonteks:\n---\n{context}\n---\n\nPertanyaan: {question}"
        )
        chain = prompt | llm
        return chain.invoke({"context": full_context[:12000], "question": question}).content

# --- Inisialisasi Session State ---
if 'step' not in st.session_state:
    st.session_state.step = "welcome"
    st.session_state.user_name = ""
    st.session_state.user_institution = ""
    st.session_state.user_email = ""
    st.session_state.document_pages = []
    st.session_state.summary = ""
    st.session_state.chat_messages = []

# --- Fungsi Render Halaman ---
def render_header():
    # --- PERUBAHAN: Membaca logo dari file lokal dan mengubahnya menjadi base64 ---
    logo_path = "SMAPER.png"
    logo_base64 = get_base64_of_bin_file(logo_path)
    
    st.markdown(f"""
    <div class="header">
        <div class="logo">
            <img src="data:image/png;base64,{logo_base64}">
            <span class="title">SMARTPAPER.AI</span>
        </div>
    </div>
    """, unsafe_allow_html=True)

def render_welcome_page():
    st.markdown("<h1>From full papers to focused insights — instantly.</h1>", unsafe_allow_html=True)
    st.markdown("<p class='tagline'>Ubah dokumen penelitian, laporan, atau artikel yang panjang menjadi ringkasan yang mudah dipahami. Dapatkan wawasan kunci dalam hitungan detik, bukan jam.</p>", unsafe_allow_html=True)
    
    with st.form("user_info_form"):
        st.markdown("<h5>Untuk memulai, silakan lengkapi data Anda:</h5>", unsafe_allow_html=True)
        name = st.text_input("Nama", placeholder="Masukkan nama lengkap Anda...")
        institution = st.text_input("Instansi", placeholder="Masukkan nama instansi atau perusahaan...")
        email = st.text_input("Alamat Email", placeholder="Masukkan alamat email Anda...")
        
        submitted = st.form_submit_button("Lanjutkan →", type="primary", use_container_width=True)
        
        if submitted and name and institution and email:
            st.session_state.user_name = name
            st.session_state.user_institution = institution
            st.session_state.user_email = email
            st.session_state.step = "input_source"
            st.rerun()
        elif submitted:
            st.error("Harap lengkapi semua kolom: Nama, Instansi, dan Alamat Email.")

def render_input_source_page():
    st.title(f"Selamat Datang, {st.session_state.user_name}!")
    st.markdown("Pilih sumber paper yang ingin Anda analisa.")
    
    source_type = st.radio("Tipe Sumber:", ("Upload File", "Cantumkan Link URL"), horizontal=True, label_visibility="collapsed")
    
    if source_type == "Upload File":
        paper_source = st.file_uploader("Pilih file (.pdf, .docx)", type=['pdf', 'docx'], label_visibility="collapsed")
    else:
        paper_source = st.text_input("Tempelkan link URL", placeholder="https://contoh.com/artikel atau link Google Drive", label_visibility="collapsed")

    if st.button("Analisa Paper Sekarang", type="primary", use_container_width=True):
        if paper_source:
            with st.spinner("AI sedang membaca dan menganalisa paper Anda..."):
                source_type_arg = 'file' if source_type == "Upload File" else 'url'
                pages = extract_text(paper_source, source_type_arg)
                
                if pages and not (isinstance(pages, list) and pages[0].startswith("Error")):
                    st.session_state.document_pages = pages
                    st.session_state.summary = generate_summary("\n".join(pages))
                    st.session_state.step = "analysis"
                    st.session_state.chat_messages = []
                    st.rerun()
                else:
                    st.error(f"Gagal memproses sumber. {pages[0] if pages else 'Tidak ada input.'}")
        else:
            st.warning("Harap sediakan sumber paper terlebih dahulu.")

def render_analysis_page():
    st.title("✅ Analisa Selesai")
    st.markdown("Berikut adalah ringkasan keseluruhan paper. Ajukan pertanyaan lebih lanjut atau minta ringkasan halaman spesifik di kolom chat bawah.")
    
    with st.container(border=True):
        st.subheader("Ringkasan Paper")
        st.markdown(st.session_state.summary)

    st.subheader("Tanya Jawab Interaktif")
    for msg in st.session_state.chat_messages:
        st.chat_message(msg["role"]).write(msg["content"])
        
    if prompt := st.chat_input("Tanya apa saja tentang paper ini... (Contoh: ringkas halaman 2)"):
        st.session_state.chat_messages.append({"role": "user", "content": prompt})
        with st.chat_message("user"): st.write(prompt)
        
        with st.chat_message("assistant"):
            with st.spinner("AI sedang berpikir..."):
                answer = answer_question(st.session_state.document_pages, prompt)
                st.session_state.chat_messages.append({"role": "assistant", "content": answer})
                st.write(answer)
    
    st.markdown("---")
    if st.button("Analisa Paper Lain", type="secondary"):
        st.session_state.step = "input_source"
        st.session_state.document_pages = []
        st.session_state.summary = ""
        st.session_state.chat_messages = []
        st.rerun()

# --- Router Utama Aplikasi ---
render_header()
col1, col2 = st.columns([0.6, 0.4], gap="large")

with col1:
    if st.session_state.step == "welcome":
        render_welcome_page()
    elif st.session_state.step == "input_source":
        render_input_source_page()
    elif st.session_state.step == "analysis":
        render_analysis_page()

with col2:
    # --- PERUBAHAN: Membaca gambar dari file lokal ---
    image_path = "paper.jfif"
    if os.path.exists(image_path):
        st.image(image_path)
